import os
from dotenv import load_dotenv
from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
import docx
import re

# Load variabel dari file .env
load_dotenv()

SECRET_KEY = os.getenv("SECRET_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
TEMPLATE_SURAT = 'template_surat_justifikasi.docx'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SECRET_KEY'] = SECRET_KEY  # Set SECRET_KEY ke Flask

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return "Tidak ada file yang dipilih!", 400
        
        file = request.files['file']

        if file.filename == '':
            return "Nama file kosong!", 400

        if file and file.filename.endswith('.docx'):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Ambil data dari dokumen PR
            data_surat = extract_info(file_path)

            if isinstance(data_surat, tuple):
                return data_surat  # Mengembalikan error jika ada

            # Buat surat justifikasi
            surat_path = generate_surat(data_surat)

            return f'''
                <h2>Surat berhasil dibuat!</h2>
                <a href="/download">Klik di sini untuk mengunduh surat</a>
            '''

    return '''
        <h2>Upload Dokumen PR (.docx)</h2>
        <form method="post" enctype="multipart/form-data">
            <input type="file" name="file">
            <input type="submit" value="Upload">
        </form>
    '''

def extract_info(file_path):
    """ Ekstrak informasi penting dari dokumen PR """
    doc = docx.Document(file_path)
    text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]  # Hilangkan baris kosong
    data = {}

    # Pola regex untuk menangkap informasi penting
    pola = {
        "NOMOR_SURAT": r"Nomor PR:\s*(.*)",
        "PERIHAL_SURAT": r"Keperluan:\s*(.*)",
        "NAMA_PENERIMA": r"Disetujui oleh:\s*(.*)",
        "ALAMAT_PENERIMA": r"Jabatan:\s*(.*)",
        "ISI_SURAT": r"Alasan Pengadaan \(Justifikasi\):\s*(.*)",
        "NAMA_PENGIRIM": r"Pemohon:\s*(.*)",
        "JABATAN_PENGIRIM": r"Jabatan:\s*(.*)"
    }

    # Ekstraksi data berdasarkan pola regex
    for key, pattern in pola.items():
        for line in text:
            match = re.search(pattern, line)
            if match:
                data[key] = match.group(1)
                break

    # Ambil rincian barang/jasa dari PR
    data["NAMA_BARANG"] = []
    data["JUMLAH_BARANG"] = []
    data["HARGA_SATUAN"] = []
    data["TOTAL_BIAYA"] = []

    barang_mode = False
    for line in text:
        if "Deskripsi Barang/Jasa:" in line:
            barang_mode = True
            continue
        if barang_mode and line.strip():
            # Contoh format: Laptop Dell XPS 15 - 5 unit
            barang_data = re.match(r"(.*) - (\d+) unit", line)
            if barang_data:
                nama_barang = barang_data.group(1).strip()
                jumlah_barang = barang_data.group(2).strip()
                data["NAMA_BARANG"].append(nama_barang)
                data["JUMLAH_BARANG"].append(jumlah_barang)
    
    harga_mode = False
    for line in text:
        if "Harga Per Unit & Total:" in line:
            harga_mode = True
            continue
        if harga_mode and line.strip():
            # Contoh format: Laptop Dell XPS 15 - Rp 25.000.000 x 5 = Rp 125.000.000
            harga_data = re.match(r".* - Rp ([\d\.]+) x (\d+) = Rp ([\d\.]+)", line)
            if harga_data:
                harga_satuan = harga_data.group(1).strip()
                total_biaya = harga_data.group(3).strip()
                data["HARGA_SATUAN"].append(harga_satuan)
                data["TOTAL_BIAYA"].append(total_biaya)

    # Konversi list menjadi string yang dipisahkan koma
    for key in ["NAMA_BARANG", "JUMLAH_BARANG", "HARGA_SATUAN", "TOTAL_BIAYA"]:
        data[key] = ", ".join(data[key]) if data[key] else "Tidak ada data"

    return data

def generate_surat(data):
    """ Mengisi template surat dengan data dari file PR """
    doc = docx.Document(TEMPLATE_SURAT)

    for para in doc.paragraphs:
        for key, value in data.items():
            if f"[{key}]" in para.text:
                para.text = para.text.replace(f"[{key}]", str(value))  # Pastikan value selalu string

    output_path = os.path.join(UPLOAD_FOLDER, "surat_output.docx")
    doc.save(output_path)
    return output_path

@app.route('/download')
def download_file():
    path = os.path.join(UPLOAD_FOLDER, "surat_output.docx")
    return send_file(path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
